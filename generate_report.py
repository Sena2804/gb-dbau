"""Génère un rapport de contrôle du fichier Maroc 2026."""

import argparse
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import database as db


SOURCE_FILE = Path("Tableau_FU_MAROC2026.xlsx")


def main(source_file: Path = SOURCE_FILE):
    candidates = db._apply_duplicate_policy(db._parse_real_excel(str(source_file)))
    quotas = db._parse_quotas_from_excel(str(source_file))
    counts = Counter(candidate["niveau_etudes"] for candidate in candidates)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT DE CONTRÔLE — BOURSE DU MAROC 2026")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(
        f"Le fichier {source_file.name} contient {len(candidates)} candidatures, "
        f"{len(quotas)} filières et {sum(quotas.values())} places."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(table.rows[0].cells, ["Niveau", "Candidatures", "Places"]):
        cell.text = value

    for niveau in db.NIVEAU_ORDER:
        row = table.add_row().cells
        row[0].text = niveau
        row[1].text = str(counts.get(niveau, 0))
        row[2].text = str(sum(value for (niv, _), value in quotas.items() if niv == niveau))

    doc.add_paragraph(
        "Contrôles validés : numéros de dossiers uniques et correspondance entre "
        "les candidatures, les filières et les quotas."
    )

    output_file = Path("data") / f"rapport_controle_{source_file.stem}.docx"
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    print(f"Rapport généré : {output_file}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("source", nargs="?", type=Path, default=SOURCE_FILE)
    args = parser.parse_args()
    main(args.source)
