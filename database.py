"""Module de gestion SQLite pour les candidatures CNBAU.
OPTIMISATIONS :
  - get_stats() → une seule requête SQL au lieu de 5
  - get_all_candidatures(), get_quotas(), get_favorables_count(), get_stats()
    sont désormais cachées via st.cache_data(ttl=2) — appelées depuis app.py
    (les fonctions ici restent pures, le cache est posé dans app.py)
"""

import io
import json
import re
import shutil
import sqlite3
import unicodedata
from datetime import datetime
from pathlib import Path

import pandas as pd

DB_PATH = "cnbau_session.db"
BACKUP_DIR = Path("backups")
GLOBAL_QUOTA_FILIERE = "Toutes filières"

TUNISIA_LICENSE_QUOTAS = {
    "sciences de l'informatique": 5,
    "obstétrique": 1,
    "economie ou gestion": 5,
    "cycle préparatoire scientifique ou technique": 2,
}
TUNISIA_MASTER_QUOTA = 4
TUNISIA_DOCTORATE_QUOTA = 1

NIVEAU_MAP = {
    "BAC + 2 ANS": "Bac + 2 ans",
    "BAC+2 ANS": "Bac + 2 ans",
    "LICENCE": "Licence",
    "MASTER": "Master",
    "DOCTORAT": "Doctorat",
    "SPECIALITE": "Spécialisation",
    "SPÉCIALISATION": "Spécialisation",
    "SPECIALISATION": "Spécialisation",
    "SPECIALITE MEDICALE": "Spécialité médicale",
    "SPÉCIALITÉ MEDICALE": "Spécialité médicale",
    "SPECIALITE MÉDICALE": "Spécialité médicale",
}

DUPLICATE_POLICY = {
    "identical": "keep",
    "same_person_diff_filiere": "keep",
    "different_people": "keep",
}

AVIS_MAP = {
    "FAVORABLE": "Favorable",
    "DEFAVORABLE": "Défavorable",
    "DÉFAVORABLE": "Défavorable",
    "SUPPLEANT": "Suppléant",
    "SUPPLÉANT": "Suppléant",
    "EN ATTENTE": "En attente",
}


def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_connection()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS travaux (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT NOT NULL,
            fichier_source TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            actif INTEGER NOT NULL DEFAULT 0
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS candidatures (
            travail_id INTEGER NOT NULL,
            id_demande TEXT NOT NULL,
            id_russe TEXT,
            numero INTEGER,
            sexe TEXT,
            name TEXT NOT NULL,
            date_lieu_naissance TEXT,
            diplome_filiere_annee TEXT,
            moyenne TEXT,
            observation TEXT,
            filiere TEXT NOT NULL,
            niveau_etudes TEXT NOT NULL,
            avis TEXT DEFAULT 'En attente',
            rang_suppleant INTEGER,
            PRIMARY KEY (travail_id, id_demande)
        )
    """)
    _migrate_multi_travaux_schema(conn)
    columns = {
        row["name"]
        for row in conn.execute("PRAGMA table_info(candidatures)").fetchall()
    }
    if "rang_suppleant" not in columns:
        conn.execute("ALTER TABLE candidatures ADD COLUMN rang_suppleant INTEGER")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS quotas (
            travail_id INTEGER NOT NULL DEFAULT 1,
            niveau_etudes TEXT NOT NULL,
            filiere TEXT NOT NULL,
            nb_places INTEGER NOT NULL,
            PRIMARY KEY (travail_id, niveau_etudes, filiere)
        )
    """)
    _normalize_all_suppleant_ranks(conn)
    conn.commit()
    conn.close()


def _table_exists(conn: sqlite3.Connection, table_name: str) -> bool:
    return conn.execute(
        "SELECT 1 FROM sqlite_master WHERE type = 'table' AND name = ?",
        (table_name,),
    ).fetchone() is not None


def _columns(conn: sqlite3.Connection, table_name: str) -> set[str]:
    if not _table_exists(conn, table_name):
        return set()
    return {row["name"] for row in conn.execute(f"PRAGMA table_info({table_name})")}


def _ensure_default_travail(conn: sqlite3.Connection) -> int:
    row = conn.execute("SELECT id FROM travaux WHERE actif = 1 ORDER BY id DESC LIMIT 1").fetchone()
    if row:
        return row["id"]
    row = conn.execute("SELECT id FROM travaux ORDER BY id LIMIT 1").fetchone()
    if row:
        conn.execute("UPDATE travaux SET actif = CASE WHEN id = ? THEN 1 ELSE 0 END", (row["id"],))
        return row["id"]
    now = datetime.now().isoformat(timespec="seconds")
    cursor = conn.execute(
        "INSERT INTO travaux (nom, fichier_source, created_at, updated_at, actif) VALUES (?, ?, ?, ?, 1)",
        ("Travail principal", "", now, now),
    )
    return cursor.lastrowid


def _migrate_multi_travaux_schema(conn: sqlite3.Connection):
    travail_id = _ensure_default_travail(conn)

    if _table_exists(conn, "candidatures"):
        candidate_columns = _columns(conn, "candidatures")
        candidate_pk = [
            row["name"]
            for row in sorted(
                conn.execute("PRAGMA table_info(candidatures)").fetchall(),
                key=lambda row: row["pk"],
            )
            if row["pk"]
        ]
    else:
        candidate_columns = set()
        candidate_pk = []

    if (
        _table_exists(conn, "candidatures")
        and (
            "travail_id" not in candidate_columns
            or candidate_pk != ["travail_id", "id_demande"]
        )
    ):
        conn.execute("ALTER TABLE candidatures RENAME TO candidatures_legacy")
        conn.execute("""
            CREATE TABLE candidatures (
                travail_id INTEGER NOT NULL,
                id_demande TEXT NOT NULL,
                id_russe TEXT,
                numero INTEGER,
                sexe TEXT,
                name TEXT NOT NULL,
                date_lieu_naissance TEXT,
                diplome_filiere_annee TEXT,
                moyenne TEXT,
                observation TEXT,
                filiere TEXT NOT NULL,
                niveau_etudes TEXT NOT NULL,
                avis TEXT DEFAULT 'En attente',
                rang_suppleant INTEGER,
                PRIMARY KEY (travail_id, id_demande)
            )
        """)
        legacy_columns = _columns(conn, "candidatures_legacy")
        select_travail = "travail_id" if "travail_id" in legacy_columns else "?"
        migration_params = () if "travail_id" in legacy_columns else (travail_id,)
        select_rank = "rang_suppleant" if "rang_suppleant" in legacy_columns else "NULL"
        conn.execute(f"""
            INSERT INTO candidatures (
                travail_id, id_demande, id_russe, numero, sexe, name,
                date_lieu_naissance, diplome_filiere_annee, moyenne, observation,
                filiere, niveau_etudes, avis, rang_suppleant
            )
            SELECT
                {select_travail}, id_demande, id_russe, numero, sexe, name,
                date_lieu_naissance, diplome_filiere_annee, moyenne, observation,
                filiere, niveau_etudes, avis, {select_rank}
            FROM candidatures_legacy
        """, migration_params)
        conn.execute("DROP TABLE candidatures_legacy")

    if _table_exists(conn, "quotas") and "travail_id" not in _columns(conn, "quotas"):
        conn.execute("ALTER TABLE quotas RENAME TO quotas_legacy")
        conn.execute("""
            CREATE TABLE quotas (
                travail_id INTEGER NOT NULL,
                niveau_etudes TEXT NOT NULL,
                filiere TEXT NOT NULL,
                nb_places INTEGER NOT NULL,
                PRIMARY KEY (travail_id, niveau_etudes, filiere)
            )
        """)
        conn.execute("""
            INSERT INTO quotas (travail_id, niveau_etudes, filiere, nb_places)
            SELECT ?, niveau_etudes, filiere, nb_places FROM quotas_legacy
        """, (travail_id,))
        conn.execute("DROP TABLE quotas_legacy")


def get_active_travail_id() -> int | None:
    if not Path(DB_PATH).exists():
        return None
    conn = get_connection()
    try:
        travail_id = _ensure_default_travail(conn)
        conn.commit()
        return travail_id
    finally:
        conn.close()


def list_travaux() -> list[dict]:
    if not Path(DB_PATH).exists():
        return []
    conn = get_connection()
    try:
        rows = conn.execute("""
            SELECT
                t.id, t.nom, t.fichier_source, t.created_at, t.updated_at, t.actif,
                COUNT(c.id_demande) AS total
            FROM travaux t
            LEFT JOIN candidatures c ON c.travail_id = t.id
            GROUP BY t.id
            ORDER BY t.updated_at DESC, t.id DESC
        """).fetchall()
        return [dict(row) for row in rows]
    except sqlite3.OperationalError:
        return []
    finally:
        conn.close()


def set_active_travail(travail_id: int):
    conn = get_connection()
    try:
        exists = conn.execute("SELECT 1 FROM travaux WHERE id = ?", (travail_id,)).fetchone()
        if not exists:
            return
        now = datetime.now().isoformat(timespec="seconds")
        conn.execute("UPDATE travaux SET actif = CASE WHEN id = ? THEN 1 ELSE 0 END", (travail_id,))
        conn.execute("UPDATE travaux SET updated_at = ? WHERE id = ?", (now, travail_id))
        conn.commit()
    finally:
        conn.close()


def _touch_active_travail(conn: sqlite3.Connection, travail_id: int | None = None):
    travail_id = travail_id or _ensure_default_travail(conn)
    now = datetime.now().isoformat(timespec="seconds")
    conn.execute("UPDATE travaux SET updated_at = ? WHERE id = ?", (now, travail_id))


def _create_travail(conn: sqlite3.Connection, nom: str, fichier_source: str = "") -> int:
    now = datetime.now().isoformat(timespec="seconds")
    conn.execute("UPDATE travaux SET actif = 0")
    cursor = conn.execute(
        "INSERT INTO travaux (nom, fichier_source, created_at, updated_at, actif) VALUES (?, ?, ?, ?, 1)",
        (nom.strip() or "Nouveau travail", fichier_source, now, now),
    )
    return cursor.lastrowid


def rename_active_travail(nom: str):
    conn = get_connection()
    try:
        travail_id = _ensure_default_travail(conn)
        now = datetime.now().isoformat(timespec="seconds")
        conn.execute(
            "UPDATE travaux SET nom = ?, updated_at = ? WHERE id = ?",
            (nom.strip() or "Travail sans nom", now, travail_id),
        )
        conn.commit()
    finally:
        conn.close()


def save_active_travail():
    conn = get_connection()
    try:
        travail_id = _ensure_default_travail(conn)
        _touch_active_travail(conn, travail_id)
        conn.commit()
    finally:
        conn.close()


def _normalize_niveau(raw: str) -> str:
    key = raw.strip().upper()
    return NIVEAU_MAP.get(key, raw.strip().title())


def _normalize_filiere(name: str) -> str:
    name = re.sub(r"^(?:fili[eè]re\s*:\s*)+", "", name, flags=re.IGNORECASE)
    return re.sub(r'\s+', ' ', name).strip()


def _normalize_header(value: object) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"\s+", " ", text).strip().upper()


def _quota_lookup_key(value: object) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"\s+", " ", text).strip().lower()


def _find_candidate_columns(ws) -> tuple[int, dict[str, int]]:
    aliases = {
        "numero": ("N°", "NO", "NUMERO"),
        "sexe": ("SEXE",),
        "id_russe": ("ID RUSSE",),
        "name": ("NOM ET PRENOM", "NOM ET PRENOMS"),
        "date_lieu_naissance": ("DATE & LIEU DE NAISSANCE", "DATE ET LIEU DE NAISSANCE"),
        "diplome_filiere_annee": ("DIPLOME / FILIERE / ANNEE",),
        "moyenne": ("MOYENNE / MENTION", "MOYENNE"),
        "observation": ("OBSERVATION",),
        "avis": ("AVIS CNABAU", "AVIS"),
    }

    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 120), values_only=True), 1):
        normalized = [_normalize_header(value) for value in row]
        if not any(value in aliases["numero"] for value in normalized):
            continue
        if not any(any(alias in value for alias in aliases["name"]) for value in normalized):
            continue

        columns = {}
        for field, field_aliases in aliases.items():
            for idx, value in enumerate(normalized):
                if any(alias in value for alias in field_aliases):
                    columns[field] = idx
                    break

        required = {"numero", "sexe", "name", "date_lieu_naissance", "diplome_filiere_annee"}
        if required.issubset(columns):
            return row_idx, columns

    raise ValueError("En-têtes de candidatures introuvables dans le fichier Excel.")


def _build_filiere_lookup() -> dict:
    quotas_path = Path(__file__).resolve().parent / "quotas.json"
    if not quotas_path.exists():
        return {}
    with open(quotas_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    lookup = {}
    for filieres in data.values():
        for fil in filieres:
            key = _normalize_filiere(fil).lower()
            lookup[key] = fil
    return lookup


def _moyenne_to_float(value: object) -> float | None:
    match = re.search(r"(\d+(?:[.,]\d+)?)", str(value or ""))
    if not match:
        return None
    try:
        return float(match.group(1).replace(",", "."))
    except ValueError:
        return None


def _session_prefix(excel_path: str, sheet_name: str | None = None) -> str:
    kind = _session_kind(excel_path, sheet_name)
    if kind == "tunisia":
        return "TUN"
    if kind == "morocco":
        return "MAR"
    return "CNA"


def _session_kind(excel_path: str, sheet_name: str | None = None) -> str:
    if sheet_name:
        normalized_sheet = _normalize_header(sheet_name)
        if normalized_sheet == "TUNISIE":
            return "tunisia"
        if normalized_sheet == "MAROC":
            return "morocco"

    stem = Path(excel_path).stem.upper()
    if "TUNIS" in stem:
        return "tunisia"
    if "MAROC" in stem:
        return "morocco"

    try:
        from openpyxl import load_workbook

        wb = load_workbook(excel_path, read_only=True, data_only=True)
        try:
            sheet_names = {_normalize_header(name) for name in wb.sheetnames}
            if "TUNISIE" in sheet_names:
                return "tunisia"
            if "MAROC" in sheet_names:
                return "morocco"
        finally:
            wb.close()
    except Exception:
        pass

    return "generic"


def _country_sheets(excel_path: str) -> list[tuple[str, str]]:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(excel_path, read_only=True, data_only=True)
        try:
            sheets = []
            for sheet_name in wb.sheetnames:
                normalized_sheet = _normalize_header(sheet_name)
                if normalized_sheet == "TUNISIE":
                    sheets.append((sheet_name, "tunisia"))
                elif normalized_sheet == "MAROC":
                    sheets.append((sheet_name, "morocco"))
            return sheets
        finally:
            wb.close()
    except Exception:
        return []


def _infer_niveau_from_diplome(diplome: str) -> str:
    text = _normalize_header(diplome)
    if "BACCALAUREAT" in text:
        return "Licence"
    if "LICENCE" in text:
        return "Master"
    if "MASTER" in text or "DEA" in text or "DIPLOME D ETUDES APPROFONDIES" in text:
        return "Doctorat"
    return "Licence"


def _looks_like_diplome(value: object) -> bool:
    text = _normalize_header(value)
    return any(
        marker in text
        for marker in (
            "BACCALAUREAT",
            "LICENCE",
            "MASTER",
            "DIPLOME",
            "DEA",
        )
    )


def _extract_niveau_from_text(text: str) -> str:
    if "NIVEAU" not in _normalize_header(text):
        return ""
    raw = text.split(":", 1)[-1].strip() if ":" in text else text
    return _normalize_niveau(raw)


def _extract_filiere_and_quota(text: str, filiere_lookup: dict) -> tuple[str, int | None] | None:
    normalized = _normalize_header(text)
    if not re.search(r"FIL+IERE\s*:", normalized):
        return None

    raw = re.sub(r"^.*?fil+i[eè]re\s*:\s*", "", text, flags=re.IGNORECASE).strip()
    quota_match = re.search(r"\(?\s*0*(\d+)\s*(?:bourses?|places?)\s*\)?", raw, flags=re.IGNORECASE)
    quota = int(quota_match.group(1)) if quota_match else None
    raw = re.sub(r"\(?\s*0*\d+\s*(?:bourses?|places?)\s*\)?", "", raw, flags=re.IGNORECASE)
    filiere = _normalize_filiere(raw)
    return filiere_lookup.get(filiere.lower(), filiere), quota


def _candidate_from_row(row) -> dict | None:
    values = [cell.value for cell in row]
    for idx, value in enumerate(values):
        if not str(value or "").strip().isdigit():
            continue
        if idx + 5 >= len(values):
            continue

        sexe = str(values[idx + 1] or "").strip().upper()
        if sexe not in {"M", "F"}:
            continue

        variants = [
            {"id_russe": "", "name": idx + 2, "date": idx + 3, "diplome": idx + 4, "moyenne": idx + 5, "observation": idx + 6, "avis": idx + 7},
            {"id_russe": str(values[idx + 2] or "").strip(), "name": idx + 3, "date": idx + 4, "diplome": idx + 5, "moyenne": idx + 6, "observation": idx + 7, "avis": idx + 8},
        ]

        selected = None
        for variant in variants:
            name = str(values[variant["name"]] or "").strip() if variant["name"] < len(values) else ""
            diplome = str(values[variant["diplome"]] or "").strip() if variant["diplome"] < len(values) else ""
            if name and _looks_like_diplome(diplome):
                selected = variant
                break

        if selected is None:
            continue

        def get_value(field: str) -> str:
            col_idx = selected[field]
            return str(values[col_idx] or "").strip() if col_idx < len(values) and values[col_idx] is not None else ""

        raw_avis = get_value("avis")
        avis = AVIS_MAP.get(raw_avis.upper(), "En attente")
        observation = get_value("observation")
        if raw_avis and raw_avis.upper() not in AVIS_MAP:
            observation = f"{observation}\nAvis initial : {raw_avis}".strip()

        return {
            "source_numero": int(str(value).strip()),
            "id_russe": selected["id_russe"],
            "sexe": sexe,
            "name": get_value("name"),
            "date_lieu_naissance": get_value("date"),
            "diplome_filiere_annee": get_value("diplome"),
            "moyenne": get_value("moyenne"),
            "observation": observation,
            "avis": avis,
        }
    return None


def _parse_real_excel_groups(excel_path: str, sheet_name: str | None = None) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(excel_path, data_only=True)
    filiere_lookup = _build_filiere_lookup()
    groups = []
    session_kind = _session_kind(excel_path, sheet_name)
    sheet_names = {_normalize_header(name) for name in wb.sheetnames}
    selected_sheet = _normalize_header(sheet_name) if sheet_name else ""
    restrict_to_tunisia_sheet = not sheet_name and session_kind == "tunisia" and "TUNISIE" in sheet_names

    def flush_group(niveau: str, filiere: str, quota: int | None, candidates: list[dict]):
        if not candidates or not filiere:
            return
        resolved_niveau = niveau or _infer_niveau_from_diplome(candidates[0].get("diplome_filiere_annee", ""))
        for candidate in candidates:
            candidate["niveau_etudes"] = resolved_niveau
            candidate["filiere"] = filiere
        candidates.sort(
            key=lambda c: _moyenne_to_float(c.get("moyenne")) if _moyenne_to_float(c.get("moyenne")) is not None else -1,
            reverse=True,
        )
        groups.append({
            "niveau_etudes": resolved_niveau,
            "filiere": filiere,
            "quota": quota if quota is not None else len(candidates),
            "candidates": candidates,
        })

    for ws in wb.worksheets:
        if selected_sheet and _normalize_header(ws.title) != selected_sheet:
            continue
        if restrict_to_tunisia_sheet and _normalize_header(ws.title) != "TUNISIE":
            continue

        current_niveau = ""
        current_filiere = ""
        current_quota = None
        current_candidates: list[dict] = []

        for row in ws.iter_rows(values_only=False):
            row_values = [cell.value for cell in row]
            row_text = " ".join(str(value).strip() for value in row_values if value not in (None, ""))
            if not row_text:
                continue

            niveau = _extract_niveau_from_text(row_text)
            if niveau:
                flush_group(current_niveau, current_filiere, current_quota, current_candidates)
                current_filiere = ""
                current_quota = None
                current_candidates = []
                current_niveau = niveau
                continue

            filiere_info = _extract_filiere_and_quota(row_text, filiere_lookup)
            if filiere_info:
                flush_group(current_niveau, current_filiere, current_quota, current_candidates)
                current_filiere, current_quota = filiere_info
                current_candidates = []
                continue

            candidate = _candidate_from_row(row)
            if candidate and current_filiere:
                current_candidates.append(candidate)

        flush_group(current_niveau, current_filiere, current_quota, current_candidates)

    wb.close()
    return groups


def _parse_real_excel(excel_path: str, sheet_name: str | None = None) -> list[dict]:
    groups = _parse_real_excel_groups(excel_path, sheet_name)
    prefix = _session_prefix(excel_path, sheet_name)
    candidates = []
    numero = 1

    for group in groups:
        for candidate in group["candidates"]:
            source_numero = candidate.pop("source_numero")
            observation = candidate.get("observation", "")
            if source_numero != numero:
                observation = f"{observation}\nN° source : {source_numero}".strip()
            candidate.update({
                "id_demande": f"{prefix}-{numero:04d}/26",
                "numero": numero,
                "observation": observation,
            })
            candidates.append(candidate)
            numero += 1
    return candidates


def _parse_quotas_from_excel(excel_path: str, sheet_name: str | None = None) -> dict:
    quotas = {}
    groups = _parse_real_excel_groups(excel_path, sheet_name)
    if _session_kind(excel_path, sheet_name) == "tunisia":
        for group in groups:
            niveau = group["niveau_etudes"]
            filiere = group["filiere"]
            if niveau == "Licence":
                quota = TUNISIA_LICENSE_QUOTAS.get(_quota_lookup_key(filiere), group["quota"])
                quotas[(niveau, filiere)] = quota
        if any(group["niveau_etudes"] == "Master" for group in groups):
            quotas[("Master", GLOBAL_QUOTA_FILIERE)] = TUNISIA_MASTER_QUOTA
        if any(group["niveau_etudes"] == "Doctorat" for group in groups):
            quotas[("Doctorat", GLOBAL_QUOTA_FILIERE)] = TUNISIA_DOCTORATE_QUOTA
        return quotas

    for group in groups:
        quotas[(group["niveau_etudes"], group["filiere"])] = group["quota"]
    return quotas


def _apply_duplicate_policy(candidates: list[dict]) -> list[dict]:
    from collections import defaultdict

    unique_by_request = {}
    for candidate in candidates:
        unique_by_request.setdefault(candidate["id_demande"], candidate)
    candidates = list(unique_by_request.values())

    by_id_russe = defaultdict(list)
    for c in candidates:
        id_russe = c.get("id_russe", "")
        if id_russe:
            by_id_russe[id_russe].append(c)

    to_drop = set()

    for id_russe, group in by_id_russe.items():
        if len(group) < 2:
            continue
        a, b = group[0], group[1]

        name_a = set(a["name"].upper().split())
        name_b = set(b["name"].upper().split())
        common = name_a & name_b
        similarity = len(common) / max(len(name_a), len(name_b)) if name_a or name_b else 0

        if similarity < 0.3:
            category = "different_people"
        elif a["filiere"] == b["filiere"] and a["niveau_etudes"] == b["niveau_etudes"]:
            category = "identical"
        else:
            category = "same_person_diff_filiere"

        if DUPLICATE_POLICY.get(category) == "drop":
            to_drop.add(b["id_demande"])

    if to_drop:
        candidates = [c for c in candidates if c["id_demande"] not in to_drop]

    return candidates


def _is_real_cnabau_file(excel_path: str) -> bool:
    from openpyxl import load_workbook

    wb = load_workbook(excel_path, read_only=True, data_only=True)
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 80), max_col=min(ws.max_column, 20), values_only=True):
                text = " ".join(str(value).strip() for value in row if value not in (None, ""))
                normalized = _normalize_header(text)
                if (
                    "NIVEAU" in normalized
                    or "BOURSE" in normalized
                    or "CNABAU" in normalized
                    or re.search(r"FIL+IERE\s*:", normalized)
                ):
                    return True
    finally:
        wb.close()
    return False


def load_excel_to_db(excel_path: str, travail_name: str | None = None) -> int:
    if _is_real_cnabau_file(excel_path):
        country_sheets = _country_sheets(excel_path)
        if len(country_sheets) > 1:
            total = 0
            first_travail_id = None
            conn = get_connection()
            try:
                for sheet_name, _ in country_sheets:
                    name = f"{travail_name or _default_travail_name(excel_path)} - {sheet_name}"
                    loaded, travail_id = _load_real_excel_into_connection(conn, excel_path, name, sheet_name)
                    if first_travail_id is None:
                        first_travail_id = travail_id
                    total += loaded
                if first_travail_id is not None:
                    conn.execute("UPDATE travaux SET actif = CASE WHEN id = ? THEN 1 ELSE 0 END", (first_travail_id,))
                conn.commit()
            finally:
                conn.close()
            return total
        if len(country_sheets) == 1:
            return _load_real_excel(excel_path, travail_name, country_sheets[0][0])
        return _load_real_excel(excel_path, travail_name)
    return _load_flat_excel(excel_path, travail_name)


def _default_travail_name(excel_path: str) -> str:
    stem = Path(excel_path).stem.replace("_", " ").replace("-", " ").strip()
    return stem.title() if stem else "Nouveau travail"


def _load_real_excel_into_connection(
    conn: sqlite3.Connection,
    excel_path: str,
    travail_name: str | None = None,
    sheet_name: str | None = None,
) -> tuple[int, int]:
    candidates = _parse_real_excel(excel_path, sheet_name)
    candidates = _apply_duplicate_policy(candidates)
    quotas = _parse_quotas_from_excel(excel_path, sheet_name)
    if not quotas:
        raise ValueError("Aucun quota n'a été trouvé dans le fichier Excel.")

    travail_id = _create_travail(conn, travail_name or _default_travail_name(excel_path), Path(excel_path).name)
    for c in candidates:
        conn.execute(
            """INSERT OR REPLACE INTO candidatures
               (travail_id, id_demande, id_russe, numero, sexe, name, date_lieu_naissance,
                diplome_filiere_annee, moyenne, observation, filiere, niveau_etudes, avis)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (travail_id, c["id_demande"], c["id_russe"], c["numero"], c["sexe"], c["name"],
             c["date_lieu_naissance"], c["diplome_filiere_annee"], c["moyenne"],
             c["observation"], c["filiere"], c["niveau_etudes"], c["avis"]),
        )
    for (niveau, filiere), nb_places in quotas.items():
        conn.execute(
            """INSERT INTO quotas (travail_id, niveau_etudes, filiere, nb_places)
               VALUES (?, ?, ?, ?)""",
            (travail_id, niveau, filiere, nb_places),
        )
    return len(candidates), travail_id


def _load_real_excel(
    excel_path: str,
    travail_name: str | None = None,
    sheet_name: str | None = None,
) -> int:
    conn = get_connection()
    try:
        total, _ = _load_real_excel_into_connection(conn, excel_path, travail_name, sheet_name)
        conn.commit()
        return total
    finally:
        conn.close()


def _load_flat_excel(excel_path: str, travail_name: str | None = None) -> int:
    df = pd.read_excel(excel_path, engine="openpyxl")
    df.columns = [c.strip() for c in df.columns]
    if "avis" in df.columns:
        df["avis"] = df["avis"].fillna("En attente").replace("", "En attente")
    else:
        df["avis"] = "En attente"

    conn = get_connection()
    travail_id = _create_travail(conn, travail_name or _default_travail_name(excel_path), Path(excel_path).name)
    for _, row in df.iterrows():
        conn.execute(
            """INSERT OR REPLACE INTO candidatures
               (travail_id, id_demande, name, filiere, niveau_etudes, avis)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (travail_id, row["id_demande"], row["name"], row["filiere"],
             row["niveau_etudes"], row["avis"]),
        )
    conn.commit()
    conn.close()
    return len(df)


def load_quotas(quotas_path: str):
    with open(quotas_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    for niveau, filieres in data.items():
        for filiere, nb_places in filieres.items():
            conn.execute(
                """INSERT OR REPLACE INTO quotas (travail_id, niveau_etudes, filiere, nb_places)
                   VALUES (?, ?, ?, ?)""",
                (travail_id, niveau, filiere, nb_places),
            )
    _touch_active_travail(conn, travail_id)
    conn.commit()
    conn.close()


def get_all_candidatures() -> pd.DataFrame:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    df = pd.read_sql_query(
        "SELECT * FROM candidatures WHERE travail_id = ?",
        conn,
        params=(travail_id,),
    )
    conn.close()
    return df


def get_quotas() -> dict:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    rows = conn.execute(
        "SELECT niveau_etudes, filiere, nb_places FROM quotas WHERE travail_id = ?",
        (travail_id,),
    ).fetchall()
    conn.close()
    return {(r["niveau_etudes"], r["filiere"]): r["nb_places"] for r in rows}


def get_favorables_count() -> dict:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    rows = conn.execute(
        """SELECT niveau_etudes, filiere, COUNT(*) as n
           FROM candidatures WHERE avis = 'Favorable' AND travail_id = ?
           GROUP BY niveau_etudes, filiere"""
        ,
        (travail_id,),
    ).fetchall()
    conn.close()
    return {(r["niveau_etudes"], r["filiere"]): r["n"] for r in rows}


def _normalize_suppleant_ranks(conn: sqlite3.Connection, niveau: str, filiere: str, travail_id: int | None = None):
    travail_id = travail_id or _ensure_default_travail(conn)
    rows = conn.execute(
        """SELECT id_demande
           FROM candidatures
           WHERE avis = 'Suppléant'
             AND travail_id = ?
             AND niveau_etudes = ?
             AND filiere = ?
           ORDER BY
             CASE WHEN rang_suppleant IS NULL THEN 1 ELSE 0 END,
             rang_suppleant,
             numero,
             id_demande""",
        (travail_id, niveau, filiere),
    ).fetchall()
    for rank, row in enumerate(rows, 1):
        conn.execute(
            "UPDATE candidatures SET rang_suppleant = ? WHERE travail_id = ? AND id_demande = ?",
            (rank, travail_id, row["id_demande"]),
        )


def _normalize_all_suppleant_ranks(conn: sqlite3.Connection):
    travail_id = _ensure_default_travail(conn)
    groups = conn.execute(
        """SELECT DISTINCT niveau_etudes, filiere
           FROM candidatures
           WHERE avis = 'Suppléant' AND travail_id = ?""",
        (travail_id,),
    ).fetchall()
    for group in groups:
        _normalize_suppleant_ranks(conn, group["niveau_etudes"], group["filiere"], travail_id)
    conn.execute(
        "UPDATE candidatures SET rang_suppleant = NULL WHERE avis != 'Suppléant' AND travail_id = ?",
        (travail_id,),
    )


def format_suppleant_rank(rank: int | None) -> str:
    try:
        rank = int(rank)
    except (TypeError, ValueError):
        return ""
    if rank <= 0:
        return ""
    return "1er suppléant" if rank == 1 else f"{rank}e suppléant"


def update_avis(id_demande: str, avis: str):
    conn = get_connection()
    try:
        travail_id = _ensure_default_travail(conn)
        candidate = conn.execute(
            """SELECT niveau_etudes, filiere, avis, rang_suppleant
               FROM candidatures WHERE travail_id = ? AND id_demande = ?""",
            (travail_id, id_demande),
        ).fetchone()
        if not candidate:
            return

        niveau = candidate["niveau_etudes"]
        filiere = candidate["filiere"]
        old_avis = candidate["avis"]

        if avis == "Suppléant" and old_avis != "Suppléant":
            next_rank = conn.execute(
                """SELECT COALESCE(MAX(rang_suppleant), 0) + 1
                   FROM candidatures
                   WHERE avis = 'Suppléant'
                     AND travail_id = ?
                     AND niveau_etudes = ?
                     AND filiere = ?""",
                (travail_id, niveau, filiere),
            ).fetchone()[0]
            conn.execute(
                """UPDATE candidatures
                   SET avis = 'Suppléant', rang_suppleant = ?
                   WHERE travail_id = ? AND id_demande = ?""",
                (next_rank, travail_id, id_demande),
            )
        elif avis == "Suppléant":
            conn.execute(
                "UPDATE candidatures SET avis = ? WHERE travail_id = ? AND id_demande = ?",
                (avis, travail_id, id_demande),
            )
        else:
            conn.execute(
                """UPDATE candidatures
                   SET avis = ?, rang_suppleant = NULL
                   WHERE travail_id = ? AND id_demande = ?""",
                (avis, travail_id, id_demande),
            )

        if old_avis == "Suppléant" or avis == "Suppléant":
            _normalize_suppleant_ranks(conn, niveau, filiere, travail_id)
        _touch_active_travail(conn, travail_id)
        conn.commit()
    finally:
        conn.close()


def search_by_field(field: str, query: str) -> dict | None:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    row = None
    if field == "numero":
        try:
            num = int(query)
            row = conn.execute(
                "SELECT * FROM candidatures WHERE travail_id = ? AND numero = ?",
                (travail_id, num),
            ).fetchone()
        except ValueError:
            pass
    elif field == "id_russe":
        row = conn.execute(
            "SELECT * FROM candidatures WHERE travail_id = ? AND id_russe = ?",
            (travail_id, query),
        ).fetchone()
    elif field == "name":
        row = conn.execute(
            "SELECT * FROM candidatures WHERE travail_id = ? AND name = ? COLLATE NOCASE",
            (travail_id, query),
        ).fetchone()
    conn.close()
    return dict(row) if row else None


def search_by_field_fuzzy(field: str, query: str) -> list[dict]:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    if field == "numero":
        rows = conn.execute(
            "SELECT * FROM candidatures WHERE travail_id = ? AND CAST(numero AS TEXT) LIKE ? ORDER BY numero LIMIT 20",
            (travail_id, f"%{query}%"),
        ).fetchall()
    elif field == "id_russe":
        rows = conn.execute(
            "SELECT * FROM candidatures WHERE travail_id = ? AND id_russe LIKE ? ORDER BY numero LIMIT 20",
            (travail_id, f"%{query}%"),
        ).fetchall()
    elif field == "name":
        rows = conn.execute(
            "SELECT * FROM candidatures WHERE travail_id = ? AND name LIKE ? COLLATE NOCASE ORDER BY numero LIMIT 20",
            (travail_id, f"%{query}%"),
        ).fetchall()
    else:
        rows = []
    conn.close()
    return [dict(r) for r in rows]


# ✅ OPTIMISATION : une seule requête SQL au lieu de 5 COUNT() séparés
def get_stats() -> dict:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    row = conn.execute("""
        SELECT
            COUNT(*) AS total,
            SUM(CASE WHEN avis = 'Favorable'   THEN 1 ELSE 0 END) AS favorables,
            SUM(CASE WHEN avis = 'Défavorable' THEN 1 ELSE 0 END) AS defavorables,
            SUM(CASE WHEN avis = 'Suppléant'   THEN 1 ELSE 0 END) AS suppleants
        FROM candidatures
        WHERE travail_id = ?
    """, (travail_id,)).fetchone()
    conn.close()
    total     = row["total"]        or 0
    fav       = row["favorables"]   or 0
    defav     = row["defavorables"] or 0
    supp      = row["suppleants"]   or 0
    traites   = fav + defav + supp
    return {
        "total":        total,
        "traites":      traites,
        "favorables":   fav,
        "defavorables": defav,
        "suppleants":   supp,
        "restants":     total - traites,
    }


NIVEAU_ORDER = ["Bac + 2 ans", "Licence", "Master", "Doctorat", "Spécialité médicale"]


def _create_base_docx():
    from itertools import groupby

    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Twips

    doc = Document()

    for section in doc.sections:
        section.top_margin = Twips(720)
        section.bottom_margin = Twips(720)
        section.left_margin = Twips(720)
        section.right_margin = Twips(720)

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Trebuchet MS"
        run.font.size = Pt(9)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)

    def set_run_font(run, size=10, bold=False, underline=False):
        run.font.name = "Trebuchet MS"
        run.font.size = Pt(size)
        run.bold = bold
        run.underline = underline

    def set_cell_shading(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    def merge_row_cells(row, table):
        n_cols = len(table.columns)
        tc = row.cells[0]._tc
        tcPr = tc.get_or_add_tcPr()
        gs = OxmlElement("w:gridSpan")
        gs.set(qn("w:val"), str(n_cols))
        tcPr.append(gs)
        tr = row._tr
        tcs = tr.findall(qn("w:tc"))
        for extra_tc in tcs[1:]:
            tr.remove(extra_tc)

    def add_table_for_section(candidates_list):
        table = doc.add_table(rows=1, cols=5, style="Table Grid")

        col_widths = [Twips(567), Twips(1418), Twips(3403), Twips(3403), Twips(2127)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

        headers = ["N° ", "RANG", "FILIERE", "NOM ET PRENOMS", "OBSERVATIONS"]
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=11)
            set_cell_shading(cell, "BFBFBF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        def sort_key(c):
            niv = c["niveau_etudes"]
            idx = NIVEAU_ORDER.index(niv) if niv in NIVEAU_ORDER else 99
            rank = c.get("rang_suppleant") or 999999
            return (idx, c["filiere"], rank, c.get("numero", 0) or 0)

        candidates_list = sorted(candidates_list, key=sort_key)

        num = 1
        for niveau, niveau_group in groupby(candidates_list, key=lambda c: c["niveau_etudes"]):
            niveau_group = list(niveau_group)
            row = table.add_row()
            cell = row.cells[0]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(niveau.upper())
            set_run_font(run, size=11, bold=True)
            set_cell_shading(cell, "E7E6E6")
            merge_row_cells(row, table)

            for filiere, fil_group in groupby(niveau_group, key=lambda c: c["filiere"]):
                fil_group = list(fil_group)
                first_row_idx = len(table.rows)

                for i_in_fil, c in enumerate(fil_group):
                    row = table.add_row()
                    cell_num = row.cells[0]
                    cell_num.text = ""
                    p = cell_num.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(str(num))
                    set_run_font(run, size=11)

                    cell_rank = row.cells[1]
                    cell_rank.text = ""
                    p = cell_rank.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(format_suppleant_rank(c.get("rang_suppleant")))
                    set_run_font(run, size=10, bold=bool(c.get("rang_suppleant")))

                    cell_fil = row.cells[2]
                    cell_fil.text = ""
                    if i_in_fil == 0:
                        p = cell_fil.paragraphs[0]
                        run = p.add_run(c["filiere"])
                        set_run_font(run, size=11)

                    cell_name = row.cells[3]
                    cell_name.text = ""
                    p = cell_name.paragraphs[0]
                    run = p.add_run(c["name"])
                    set_run_font(run, size=11)

                    cell_obs = row.cells[4]
                    cell_obs.text = ""
                    p = cell_obs.paragraphs[0]
                    obs = c.get("observation") or ""
                    run = p.add_run(obs)
                    set_run_font(run, size=11)

                    num += 1

                last_row_idx = len(table.rows) - 1
                if last_row_idx > first_row_idx:
                    table.cell(first_row_idx, 2).merge(table.cell(last_row_idx, 2))

        return table

    logo_path = Path(__file__).parent / "assets" / "logo.png"
    if logo_path.exists():
        section = doc.sections[0]
        avail_width = section.page_width - section.left_margin - section.right_margin
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo_path), width=avail_width)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DIRECTION DES BOURSES ET AIDES UNIVERSITAIRES ")
    set_run_font(run, size=10, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "LISTE DES ETUDIANTS PRESELECTIONNES POUR BENEFICIER DE LA BOURSE "
        "DE COOPERATION TUNISIENNE AU TITRE DE L\u2019ANNEE ACADEMIQUE 2026-2027"
    )
    set_run_font(run, size=10, bold=True)

    doc.add_paragraph()

    return doc, add_table_for_section, set_run_font


def export_to_docx(output_path: str) -> str:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    rows = conn.execute(
        """SELECT * FROM candidatures
           WHERE travail_id = ? AND avis IN ('Favorable', 'Suppléant')
           ORDER BY niveau_etudes, filiere,
                    CASE WHEN avis = 'Suppléant' THEN rang_suppleant ELSE numero END,
                    numero""",
        (travail_id,),
    ).fetchall()
    conn.close()
    candidates = [dict(r) for r in rows]

    favorables = [c for c in candidates if c["avis"] == "Favorable"]
    suppleants = [c for c in candidates if c["avis"] == "Suppléant"]

    doc, add_table_for_section, set_run_font = _create_base_docx()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("LISTE DES CANDIDATS TITULAIRES ")
    set_run_font(run, size=13, underline=True)

    add_table_for_section(favorables)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("LISTE DES CANDIDATS SUPPLÉANTS")
    set_run_font(run, size=13, underline=True)

    add_table_for_section(suppleants)

    doc.save(output_path)
    return output_path


def export_all_avis_to_docx(output_path: str) -> str:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    rows = conn.execute(
        """SELECT * FROM candidatures
           WHERE travail_id = ? AND avis IN ('Favorable', 'Suppléant', 'Défavorable')
           ORDER BY niveau_etudes, filiere,
                    CASE WHEN avis = 'Suppléant' THEN rang_suppleant ELSE numero END,
                    numero""",
        (travail_id,),
    ).fetchall()
    conn.close()
    candidates = [dict(r) for r in rows]

    favorables  = [c for c in candidates if c["avis"] == "Favorable"]
    suppleants  = [c for c in candidates if c["avis"] == "Suppléant"]
    defavorables = [c for c in candidates if c["avis"] == "Défavorable"]

    doc, add_table_for_section, set_run_font = _create_base_docx()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("LISTE DES CANDIDATS TITULAIRES ")
    set_run_font(run, size=13, underline=True)
    add_table_for_section(favorables)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("LISTE DES CANDIDATS SUPPLÉANTS")
    set_run_font(run, size=13, underline=True)
    add_table_for_section(suppleants)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("LISTE DES CANDIDATS NON RETENUS")
    set_run_font(run, size=13, underline=True)
    add_table_for_section(defavorables)

    doc.save(output_path)
    return output_path


def export_avis_to_xlsx(output_path: str) -> str:
    from itertools import groupby

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    conn = get_connection()
    travail_id = _ensure_default_travail(conn)

    avis_config = [
        ("Favorable",  "Favorables (Titulaires)"),
        ("Suppléant",  "Suppléants"),
        ("Défavorable","Défavorables"),
    ]

    wb = Workbook()
    wb.remove(wb.active)

    header_font  = Font(name="Trebuchet MS", bold=True, size=11)
    header_fill  = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")
    niveau_fill  = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
    headers      = ["N°", "Rang suppléant", "Filière", "Nom et Prénoms", "Observation"]
    col_widths   = [8, 18, 45, 35, 25]

    for avis_value, sheet_name in avis_config:
        rows = conn.execute(
            """SELECT numero, name, filiere, niveau_etudes, observation, rang_suppleant
               FROM candidatures
               WHERE travail_id = ? AND avis = ?
               ORDER BY niveau_etudes, filiere,
                        CASE WHEN avis = 'Suppléant' THEN rang_suppleant ELSE numero END,
                        numero""",
            (travail_id, avis_value),
        ).fetchall()

        ws = wb.create_sheet(title=sheet_name)

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        for col_idx, width in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = width

        candidates = [dict(r) for r in rows]

        def sort_key(c):
            niv = c["niveau_etudes"]
            idx = NIVEAU_ORDER.index(niv) if niv in NIVEAU_ORDER else 99
            rank = c.get("rang_suppleant") or 999999
            return (idx, c["filiere"], rank, c.get("numero", 0) or 0)

        candidates.sort(key=sort_key)

        current_row = 2
        for niveau, niveau_group in groupby(candidates, key=lambda c: c["niveau_etudes"]):
            cell = ws.cell(row=current_row, column=1, value=niveau.upper())
            cell.font = Font(name="Trebuchet MS", bold=True, size=11)
            for col_idx in range(1, len(headers) + 1):
                ws.cell(row=current_row, column=col_idx).fill = niveau_fill
            ws.merge_cells(
                start_row=current_row, start_column=1,
                end_row=current_row, end_column=len(headers),
            )
            current_row += 1

            for c in niveau_group:
                ws.cell(row=current_row, column=1, value=c.get("numero", ""))
                ws.cell(
                    row=current_row,
                    column=2,
                    value=format_suppleant_rank(c.get("rang_suppleant")),
                )
                ws.cell(row=current_row, column=3, value=c.get("filiere", ""))
                ws.cell(row=current_row, column=4, value=c.get("name", ""))
                ws.cell(row=current_row, column=5, value=c.get("observation", ""))
                current_row += 1

    conn.close()
    wb.save(output_path)
    return output_path


def export_quotas_to_xlsx(output_path: str) -> str:
    from itertools import groupby

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    quotas_rows = conn.execute(
        "SELECT niveau_etudes, filiere, nb_places FROM quotas WHERE travail_id = ? ORDER BY niveau_etudes, filiere",
        (travail_id,),
    ).fetchall()

    fav_rows = conn.execute(
        """SELECT niveau_etudes, filiere, COUNT(*) as n
           FROM candidatures WHERE travail_id = ? AND avis = 'Favorable'
           GROUP BY niveau_etudes, filiere""",
        (travail_id,),
    ).fetchall()
    conn.close()

    fav_counts = {(r["niveau_etudes"], r["filiere"]): r["n"] for r in fav_rows}

    wb = Workbook()
    ws = wb.active
    ws.title = "Quotas par Filière"

    headers     = ["Niveau", "Filière", "Places (Quota)", "Favorables", "Restantes"]
    header_font = Font(name="Trebuchet MS", bold=True, size=11)
    header_fill = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")
    niveau_fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
    col_widths  = [18, 50, 16, 14, 14]

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for col_idx, width in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    quotas_data = [dict(r) for r in quotas_rows]

    def sort_key(q):
        niv = q["niveau_etudes"]
        idx = NIVEAU_ORDER.index(niv) if niv in NIVEAU_ORDER else 99
        return (idx, q["filiere"])

    quotas_data.sort(key=sort_key)

    current_row    = 2
    total_places   = 0
    total_fav      = 0
    total_restantes = 0

    for niveau, group in groupby(quotas_data, key=lambda q: q["niveau_etudes"]):
        cell = ws.cell(row=current_row, column=1, value=niveau.upper())
        cell.font = Font(name="Trebuchet MS", bold=True, size=12)
        for col_idx in range(1, len(headers) + 1):
            ws.cell(row=current_row, column=col_idx).fill = niveau_fill
        ws.merge_cells(
            start_row=current_row, start_column=1,
            end_row=current_row, end_column=len(headers),
        )
        current_row += 1

        for q in group:
            key       = (q["niveau_etudes"], q["filiere"])
            if q["filiere"] == GLOBAL_QUOTA_FILIERE:
                fav = sum(count for (niv, _), count in fav_counts.items() if niv == q["niveau_etudes"])
            else:
                fav = fav_counts.get(key, 0)
            restantes = q["nb_places"] - fav

            total_places    += q["nb_places"]
            total_fav       += fav
            total_restantes += restantes

            ws.cell(row=current_row, column=1, value=q["niveau_etudes"])
            ws.cell(row=current_row, column=2, value=q["filiere"])
            ws.cell(row=current_row, column=3, value=q["nb_places"]).alignment = Alignment(horizontal="center")
            ws.cell(row=current_row, column=4, value=fav).alignment             = Alignment(horizontal="center")
            ws.cell(row=current_row, column=5, value=restantes).alignment       = Alignment(horizontal="center")
            current_row += 1

    total_fill = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")
    total_font = Font(name="Trebuchet MS", bold=True, size=11)
    ws.cell(row=current_row, column=2, value="TOTAL").font = total_font
    for col_idx in range(1, len(headers) + 1):
        ws.cell(row=current_row, column=col_idx).fill = total_fill
    ws.cell(row=current_row, column=3, value=total_places).font  = total_font
    ws.cell(row=current_row, column=3).alignment = Alignment(horizontal="center")
    ws.cell(row=current_row, column=4, value=total_fav).font     = total_font
    ws.cell(row=current_row, column=4).alignment = Alignment(horizontal="center")
    ws.cell(row=current_row, column=5, value=total_restantes).font = total_font
    ws.cell(row=current_row, column=5).alignment = Alignment(horizontal="center")

    wb.save(output_path)
    return output_path


def get_total_quota() -> int:
    conn = get_connection()
    travail_id = _ensure_default_travail(conn)
    total = conn.execute(
        "SELECT COALESCE(SUM(nb_places), 0) FROM quotas WHERE travail_id = ?",
        (travail_id,),
    ).fetchone()[0]
    conn.close()
    return total


def transfer_quota(source_niveau: str, source_filiere: str,
                   dest_niveau: str, dest_filiere: str,
                   nb_places: int) -> dict:
    if nb_places <= 0:
        return {"success": False, "error": "Le nombre de places doit être supérieur à 0."}

    if source_niveau == dest_niveau and source_filiere == dest_filiere:
        return {"success": False, "error": "La source et la destination doivent être différentes."}

    conn = get_connection()
    try:
        travail_id = _ensure_default_travail(conn)
        row_src = conn.execute(
            "SELECT nb_places FROM quotas WHERE travail_id = ? AND niveau_etudes = ? AND filiere = ?",
            (travail_id, source_niveau, source_filiere),
        ).fetchone()
        if not row_src:
            return {"success": False, "error": f"Quota source introuvable ({source_niveau}, {source_filiere})."}

        quota_source = row_src["nb_places"]

        if source_filiere == GLOBAL_QUOTA_FILIERE:
            fav_row = conn.execute(
                "SELECT COUNT(*) as n FROM candidatures WHERE travail_id = ? AND avis = 'Favorable' AND niveau_etudes = ?",
                (travail_id, source_niveau),
            ).fetchone()
        else:
            fav_row = conn.execute(
                "SELECT COUNT(*) as n FROM candidatures WHERE travail_id = ? AND avis = 'Favorable' AND niveau_etudes = ? AND filiere = ?",
                (travail_id, source_niveau, source_filiere),
            ).fetchone()
        fav_source  = fav_row["n"]
        disponibles = quota_source - fav_source

        if nb_places > disponibles:
            return {
                "success": False,
                "error": f"Places disponibles insuffisantes. Quota : {quota_source}, Favorables : {fav_source}, Disponibles : {disponibles}.",
            }

        row_dest = conn.execute(
            "SELECT nb_places FROM quotas WHERE travail_id = ? AND niveau_etudes = ? AND filiere = ?",
            (travail_id, dest_niveau, dest_filiere),
        ).fetchone()
        if not row_dest:
            return {"success": False, "error": f"Quota destination introuvable ({dest_niveau}, {dest_filiere})."}

        quota_dest = row_dest["nb_places"]

        conn.execute(
            "UPDATE quotas SET nb_places = nb_places - ? WHERE travail_id = ? AND niveau_etudes = ? AND filiere = ?",
            (nb_places, travail_id, source_niveau, source_filiere),
        )
        conn.execute(
            "UPDATE quotas SET nb_places = nb_places + ? WHERE travail_id = ? AND niveau_etudes = ? AND filiere = ?",
            (nb_places, travail_id, dest_niveau, dest_filiere),
        )
        _touch_active_travail(conn, travail_id)
        conn.commit()

        return {
            "success": True,
            "source_nouveau": quota_source - nb_places,
            "dest_nouveau":   quota_dest   + nb_places,
        }
    except Exception as e:
        conn.rollback()
        return {"success": False, "error": str(e)}
    finally:
        conn.close()


def is_db_loaded() -> bool:
    if not Path(DB_PATH).exists():
        return False
    conn = get_connection()
    try:
        travail_id = _ensure_default_travail(conn)
        count = conn.execute(
            "SELECT COUNT(*) FROM candidatures WHERE travail_id = ?",
            (travail_id,),
        ).fetchone()[0]
    except sqlite3.OperationalError:
        conn.close()
        return False
    conn.close()
    return count > 0


def backup_db(label: str = "manual") -> Path:
    source_path = Path(DB_PATH)
    if not source_path.exists():
        raise FileNotFoundError("Aucune base de donnees a sauvegarder.")

    BACKUP_DIR.mkdir(exist_ok=True)
    safe_label = re.sub(r"[^A-Za-z0-9_-]+", "_", label).strip("_") or "manual"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"cnbau_session_{safe_label}_{timestamp}.db"

    source = sqlite3.connect(source_path)
    target = sqlite3.connect(backup_path)
    try:
        source.backup(target)
    finally:
        target.close()
        source.close()

    return backup_path


def read_db_bytes() -> bytes:
    db_path = Path(DB_PATH)
    if not db_path.exists():
        return b""
    return db_path.read_bytes()


def summarize_db(path: str | Path = DB_PATH) -> dict:
    db_path = Path(path)
    if not db_path.exists():
        return {"exists": False, "travaux": 0, "candidatures": 0, "avis": {}}

    conn = sqlite3.connect(db_path)
    try:
        conn.row_factory = sqlite3.Row
        tables = {
            row["name"]
            for row in conn.execute(
                "SELECT name FROM sqlite_master WHERE type = 'table'"
            ).fetchall()
        }
        if not {"travaux", "candidatures", "quotas"}.issubset(tables):
            raise ValueError("Le fichier ne contient pas une sauvegarde CNBAU valide.")

        travaux = conn.execute("SELECT COUNT(*) FROM travaux").fetchone()[0]
        candidatures = conn.execute("SELECT COUNT(*) FROM candidatures").fetchone()[0]
        avis_rows = conn.execute(
            "SELECT avis, COUNT(*) AS total FROM candidatures GROUP BY avis ORDER BY avis"
        ).fetchall()
        avis = {row["avis"]: row["total"] for row in avis_rows}
        return {
            "exists": True,
            "travaux": travaux,
            "candidatures": candidatures,
            "avis": avis,
        }
    finally:
        conn.close()


def restore_db_from_bytes(data: bytes) -> dict:
    if not data:
        raise ValueError("La sauvegarde est vide.")

    temp_path = Path(f"{DB_PATH}.restore.tmp")
    temp_path.write_bytes(data)

    try:
        conn = sqlite3.connect(temp_path)
        try:
            integrity = conn.execute("PRAGMA integrity_check").fetchone()[0]
            if integrity != "ok":
                raise ValueError(f"Sauvegarde SQLite invalide : {integrity}")
        finally:
            conn.close()

        summary = summarize_db(temp_path)
        if Path(DB_PATH).exists():
            backup_db("before_restore")
        shutil.move(str(temp_path), DB_PATH)
        return summary
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise


def reset_db():
    if Path(DB_PATH).exists():
        backup_db("before_reset")
        Path(DB_PATH).unlink()
